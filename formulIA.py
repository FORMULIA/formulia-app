import streamlit as st
from openpyxl import load_workbook

st.set_page_config(page_title="FormulIA - Generador de propuestas", layout="centered")

st.title("🧠 FormulIA")
st.subheader("Sistema automatizado para estructuración de propuestas")
st.markdown("---")

# Paso 1: Componentes

st.header("1️⃣ ¿Qué componentes incluye tu propuesta?")
componentes = st.multiselect(
    "Selecciona uno o varios:",
    ["Formación", "Monitoreo y Evaluación", "Materiales", "Operación"]
)

# Paso 2: Modalidad
modalidades = {}
for comp in ["Formación", "Operación"]:
    if comp in componentes:
        modalidades[comp] = st.selectbox(
            f"¿Qué modalidad tendrá el componente {comp}?",
            ["Presencial", "Virtual", "Híbrida"],
            key=f"modalidad_{comp}"
        )

# Paso 3: Organización y municipio
st.header("2️⃣ Datos generales del proyecto")

# Mostrar los campos con valor por defecto desde session_state
org = st.text_input(
    "¿A qué organización se dirige la propuesta?",
    value=st.session_state.get("organizacion", "")
)
municipio = st.text_input(
    "¿En qué municipio se ejecutará el proyecto?",
    value=st.session_state.get("municipio", "")
)

# Guardar los valores en session_state para uso posterior
if org:
    st.session_state["organizacion"] = org
if municipio:
    st.session_state["municipio"] = municipio


# Paso 4: Estrategias pedagógicas
st.header("3️⃣ Estrategias pedagógicas")
estrategias = st.multiselect(
    "¿Qué estrategias deseas aplicar?",
    ["Transición", "Primero", "Remediación"]
)

# Paso 5: Sedes educativas
st.header("4️⃣ Sedes educativas")

# Guardar número de sedes
num_sedes = st.number_input(
    "¿Cuántas sedes serán impactadas?",
    min_value=1,
    value=st.session_state.get("num_sedes", 1),
    step=1
)
st.session_state["num_sedes"] = num_sedes

# Capturar nombres de sedes y guardarlos
sedes = st.session_state.get("sedes", [""] * int(num_sedes))

for i in range(int(num_sedes)):
    sede = st.text_input(f"Nombre de la sede #{i+1}", value=sedes[i], key=f"sede_{i}")
    sedes[i] = sede

st.session_state["sedes"] = sedes

# Paso 6: Estudiantes y docentes
st.header("5️⃣ Registro de estudiantes y docentes por sede")

poblacion_por_sede = {}
grados_por_estrategia = {
    "Transición": ["0°"],
    "Primero": ["1°"],
    "Remediación": ["2°", "3°", "4°", "5°"]
}

grados_requeridos = set()
for est in estrategias:
    grados_requeridos.update(grados_por_estrategia[est])

incluir_todos_docentes = st.radio(
    "¿Deseas incluir a los docentes de todos los grados?",
    ["Sí", "No"]
)

for sede in sedes:
    st.subheader(f"Sede: {sede}")
    poblacion_por_sede[sede] = {}
    for grado in ["0°", "1°", "2°", "3°", "4°", "5°"]:
        incluir_grado = grado in grados_requeridos or incluir_todos_docentes == "Sí"
        if incluir_grado:
            col1, col2 = st.columns(2)
            with col1:
                est = st.number_input(f"👩‍🎓 Estudiantes en {grado} - {sede}", min_value=0, key=f"{sede}_{grado}_est")
            with col2:
                doc = st.number_input(f"👨‍🏫 Docentes en {grado} - {sede}", min_value=0, key=f"{sede}_{grado}_doc")
            if doc == 0 and est > 0:
                doc = round(est / 25)
                st.markdown(f"🔎 *Docentes estimados: {doc}*")
            poblacion_por_sede[sede][grado] = {"estudiantes": est, "docentes": doc}

st.session_state["poblacion_por_sede"] = poblacion_por_sede

# Calcular totales globales
total_estudiantes = 0
total_docentes = 0

for sede_data in st.session_state["poblacion_por_sede"].values():
    for grado_data in sede_data.values():
        total_estudiantes += grado_data["estudiantes"]
        total_docentes += grado_data["docentes"]

# Guardar en session_state
st.session_state["total_estudiantes"] = total_estudiantes
st.session_state["total_docentes"] = total_docentes

# Paso 7: Materiales
st.header("6️⃣ Selección de materiales")
materiales_por_estrategia = {
    "Transición": ["Cartilla docente transición", "Cartilla estudiante transición", "Cartilla de cuentos transición", "Kit aula transición"],
    "Primero": ["Guía docente tomo I", "Guía docente tomo II", "Guía estudiante unidad I", "Guía estudiante unidad II", "Guía estudiante unidad III", "Guía estudiante unidad IV", "Libro de cuentos", "Big Book", "Fichas", "Componedores aula", "Componedores individuales"],
    "Remediación": ["Guía docente remediación", "Guía estudiante remediación", "Cartilla cuentos remediación", "Fichas de apoyo remediación"]
}
materiales_seleccionados = {}
for estrategia in estrategias:
    st.subheader(f"📦 Materiales para {estrategia}")
    seleccion = st.multiselect(
        f"Selecciona los materiales que deseas incluir para {estrategia}",
        materiales_por_estrategia[estrategia],
        key=f"materiales_{estrategia}"
    )
    materiales_seleccionados[estrategia] = seleccion

st.session_state["materiales_seleccionados"] = materiales_seleccionados

# Paso 8: Logística (si presencial)
if "Formación" in componentes and modalidades.get("Formación") == "Presencial":
    st.header("7️⃣ Logística de formación (Presencial)")

    # Inicializa diccionario si no existe
    if "formacion_logistica" not in st.session_state:
        st.session_state["formacion_logistica"] = {}

    # Número de viajes estimados
    num_viajes = st.number_input(
        "¿Cuántos viajes estimas para desarrollar las formaciones?",
        min_value=1,
        max_value=10,
        value=3,
        step=1
    )
    st.session_state["formacion_logistica"]["num_viajes"] = num_viajes

    # Horas estimadas de viaje ida y vuelta
    horas_viaje = st.number_input(
        "¿Cuántas horas estimas de desplazamiento ida y vuelta?",
        min_value=1,
        max_value=24,
        value=4,
        step=1
    )
    st.session_state["formacion_logistica"]["horas_viaje"] = horas_viaje

    # Tipo de transporte
    tipo_transporte = st.radio("¿Qué tipo de transporte se usará?", ["Terrestre", "Aéreo"])
    st.session_state["formacion_logistica"]["tipo_transporte"] = tipo_transporte

    # Costo promedio de transporte por visita
    costo_transporte = st.number_input(
        "Costo promedio por visita (COP)",
        min_value=100000,
        value=500000 if tipo_transporte == "Aéreo" else 150000,
        step=50000
    )
    st.session_state["formacion_logistica"]["costo_transporte"] = costo_transporte

    # Costo promedio de hotel por noche
    valor_hotel = st.number_input(
        "Costo promedio por noche de hotel (COP)",
        min_value=50000,
        value=180000,
        step=10000
    )
    st.session_state["formacion_logistica"]["valor_hotel"] = valor_hotel

    # Cargar temas desde Excel (B3:B9)
    try:
        wb_temas = load_workbook("estructura de costos formuLIA.xlsx", data_only=True)
        ws_temas = wb_temas[" FORMACIÓN"]
        temas_formacion = [ws_temas[f"B{row}"].value for row in range(3, 10) if ws_temas[f"B{row}"].value]
    except Exception as e:
        st.error(f"❌ No se pudieron cargar los temas desde el Excel: {e}")
        temas_formacion = []

    temas_seleccionados = st.multiselect(
        "📚 Selecciona los temas que deseas trabajar:",
        temas_formacion
    )
    st.session_state["formacion_logistica"]["temas"] = temas_seleccionados

    # BLOQUE DE REFRIGERIOS
    incluir_refrigerios = st.radio("¿Deseas incluir refrigerios?", ["Sí", "No"])
    if incluir_refrigerios == "Sí":
        try:
            wb_temp = load_workbook("estructura de costos formuLIA.xlsx", data_only=True)
            ws_temp = wb_temp[" FORMACIÓN"]
            valor_unitario = ws_temp["C24"].value or 8000
        except:
            valor_unitario = 8000

        st.markdown(f"💰 Valor actual del refrigerio (desde Excel): **${int(valor_unitario):,} COP**")

        actualizar_valor = st.radio("¿Deseas actualizar el valor del refrigerio?", ["No", "Sí"])
        if actualizar_valor == "Sí":
            valor_unitario = st.number_input(
                "Nuevo valor del refrigerio (COP)",
                min_value=1000,
                value=int(valor_unitario),
                step=500
            )

        num_sesiones = st.number_input("¿En cuántas sesiones se ofrecerán refrigerios?", min_value=1, step=1)

        # Total docentes desde el estado (sumando todas las sedes)
        total_docentes = sum(info["docentes"] for sede in st.session_state["poblacion_por_sede"].values() for info in sede.values())
        cantidad_refrigerios = int(round(total_docentes * 1.2 * num_sesiones))

        st.success(f"🥤 Total estimado de refrigerios: {cantidad_refrigerios}")

        st.session_state["refrigerios"] = {
            "valor_unitario": valor_unitario,
            "num_sesiones": num_sesiones,
            "cantidad_refrigerios": cantidad_refrigerios
        }
# Paso 9: Grupos de formación
if "Formación" in componentes:
    st.header("8️⃣ Grupos de formación")
    total_docentes = sum(info["docentes"] for sede in poblacion_por_sede.values() for info in sede.values())
    n_grupos = (total_docentes + 39) // 40
    st.success(f"👥 Total de docentes: {total_docentes}")
    st.info(f"🔹 Se requerirán {n_grupos} grupo(s) (máx. 40 personas por grupo)")
    st.session_state["grupos_formacion"] = {
        "total_docentes": total_docentes,
        "n_grupos": n_grupos
    }

# Paso 10: Remediación
if "Remediación" in estrategias:
    st.header("9️⃣ Estrategia de Remediación")
    porcentaje_remediacion = st.slider("¿Qué porcentaje de estudiantes de 2° a 5° requieren remediación?", 0, 100, 25)
    total_estudiantes_remediacion = sum(
        info["estudiantes"]
        for sede in poblacion_por_sede.values()
        for grado, info in sede.items()
        if grado in ["2°", "3°", "4°", "5°"]
    )
    estimado = int(round(total_estudiantes_remediacion * porcentaje_remediacion / 100))
    st.success(f"👧🧒 Se estima que {estimado} estudiantes requieren remediación")
    st.session_state["remediacion"] = {
        "porcentaje": porcentaje_remediacion,
        "total_grados_2_5": total_estudiantes_remediacion,
        "estimado_remediacion": estimado
    }

# Paso 11: Selección de pruebas (Monitoreo)
if "Monitoreo y Evaluación" in componentes and any(e in estrategias for e in ["Primero", "Remediación"]):
    st.header("🔍 Módulo de Monitoreo y Evaluación")
    st.markdown("### 📁 Selecciona las pruebas que deseas aplicar:")
    pruebas_disponibles = ["EGRA entrada", "EGRA salida"]
    if "Primero" in estrategias:
        pruebas_disponibles.extend(["Semana 10", "Semana 20", "Semana 30", "Semana 40"])
    if "Remediación" in estrategias:
        pruebas_disponibles.extend(["Semana 1", "Semana 7", "Semana 14"])
    pruebas_seleccionadas = st.multiselect("Pruebas disponibles según las estrategias seleccionadas:", pruebas_disponibles)
    st.session_state["pruebas_monitoreo"] = pruebas_seleccionadas

# Paso 12: Aplicación de pruebas y productos asociados
if "pruebas_monitoreo" in st.session_state and st.session_state["pruebas_monitoreo"]:
    st.header("📊 Aplicación de pruebas y productos por evaluación")

    productos_estandar = [
        "Informe de resultados",
        "Presentación a actores",
        "Informe técnico consolidado",
        "Devolución a instituciones"
    ]

    configuracion_pruebas = {}

    for prueba in st.session_state["pruebas_monitoreo"]:
        st.subheader(f"🧪 {prueba}")

        incluir_aplicacion = st.radio(
            f"¿Deseas incluir la aplicación de la prueba '{prueba}'?",
            ["Sí", "No"],
            key=f"aplicacion_{prueba}"
        )

        productos_incluidos = st.multiselect(
            f"Selecciona los productos que deseas incluir para '{prueba}':",
            productos_estandar,
            default=productos_estandar,
            key=f"productos_{prueba}"
        )

        configuracion_pruebas[prueba] = {
            "aplicacion": incluir_aplicacion == "Sí",
            "productos": productos_incluidos
        }

    st.session_state["configuracion_pruebas"] = configuracion_pruebas

from openpyxl import load_workbook
from io import BytesIO

st.markdown("---")
st.subheader("📤 Exportar archivo Excel")

# Pregunta AIU antes del botón de exportación
st.markdown("---")
st.subheader("⚙️ Porcentaje de AIU")
aiu = st.number_input(
    "Indica el % de AIU (Administración, Imprevistos y Utilidad):",
    min_value=0,
    max_value=100,
    value=35,
    step=1
)
st.session_state["aiu_porcentaje"] = aiu

# Botón para exportar
if st.button("📥 Generar archivo Excel con datos"):
    excel_path = "estructura de costos formuLIA.xlsx"

    try:
        wb = load_workbook(excel_path)
        ws = wb[" FORMACIÓN"]

        # Recuperar datos
        formacion = st.session_state.get("formacion_logistica", {})
        grupos = st.session_state.get("grupos_formacion", {}).get("n_grupos", 1)
        num_viajes = formacion.get("num_viajes", 3)
        horas_viaje = formacion.get("horas_viaje", 0)
        temas = formacion.get("temas", [])

        costo_transporte = formacion.get("costo_transporte", 0)
        valor_hotel = formacion.get("valor_hotel", 0)

        refrigerios = st.session_state.get("refrigerios", {})
        valor_unitario_refrigerio = refrigerios.get("valor_unitario", 8000)
        cantidad_refrigerios = refrigerios.get("cantidad_refrigerios", 0)
        num_sesiones = refrigerios.get("num_sesiones", 0)

        total_docentes = sum(
            info["docentes"]
            for sede in st.session_state["poblacion_por_sede"].values()
            for info in sede.values()
        )

        aiu = st.session_state.get("aiu_porcentaje", 35)

        # ACTUALIZAR HOJA DE FORMACIÓN
        for row in range(3, 10):
            tema = ws[f"B{row}"].value
            if tema in temas:
                horas = ws[f"C{row}"].value
                if isinstance(horas, (int, float)):
                    ws[f"C{row}"] = horas * grupos
                ws[f"F{row}"] = horas_viaje
                ws[f"J{row}"] = aiu / 100
                # No tocamos O si está seleccionado
            else:
                ws[f"C{row}"] = 0
                ws[f"F{row}"] = 0
                ws[f"J{row}"] = aiu / 100
                ws[f"O{row}"] = 0  # ← limpiar columna O si el tema no fue seleccionado

        # Otras celdas clave
        ws["C14"] = valor_hotel
        ws["C16"] = costo_transporte
        ws["D16"] = num_viajes
        ws["D14"] = num_viajes * 3
        ws["D15"] = num_viajes * 3
        ws["C24"] = valor_unitario_refrigerio
        ws["C23"] = int(round(total_docentes * 1.2))
        ws["C25"] = num_sesiones
        ws["E33"] = total_docentes
        ws["E18"] = aiu / 100
        ws["C27"] = aiu / 100

        # Guardar y descargar
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        st.success("✅ Archivo Excel actualizado correctamente.")
        st.download_button(
            label="⬇️ Descargar archivo actualizado",
            data=output,
            file_name="formulIA_actualizado.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        st.error(f"❌ Error al procesar el archivo Excel: {e}")

from docx import Document
import streamlit as st

# === Cargar documento base ===
ruta_word = "Ejemplo propuesta.docx"
ruta_salida = "Propuesta_FormulIA.docx"
doc = Document(ruta_word)

# === Extraer datos desde session_state ===
nombre_organizacion = st.session_state.get("organizacion", "ORGANIZACIÓN")
municipio = st.session_state.get("municipio", "MUNICIPIO")
num_sedes = st.session_state.get("num_sedes", 0)
lista_sedes = st.session_state.get("sedes", [])

# === Función para pluralizar correctamente ===
def frase_instituciones(n):
    return "1 Institución Educativa" if n == 1 else f"{n} Instituciones Educativas"

texto_sedes = frase_instituciones(num_sedes)

# === Preparar texto de sedes con gramática correcta ===
if num_sedes == 1 and lista_sedes:
    sedes_como_texto = f"Institución Educativa {lista_sedes[0]}"
else:
    sedes_como_texto = ", ".join([f"Institución Educativa {s}" for s in lista_sedes if s])

# === Lista de frases posibles que indican cantidad de sedes ===
reemplazos_cantidad = [
    "3 Instituciones Educativas",
    "3 instituciones educativas",
    "3 sedes",
    "3 I.E.",
    "tres instituciones educativas",
    "Tres Instituciones Educativas"
]

# === Frase con nombres fijos en la viñeta de población focalizada ===
marcador_sedes_fijas = "Luis Felipe Cabrera, Institución Educativa de Santa Ana, Institución Educativa De Ararca"

# === Viñetas antiguas que deben ser eliminadas ===
viñetas_antiguas = [
    "• 37 docentes de transición a tercero",
    "• 568 estudiantes de transición a primero",
    "• 525 estudiantes de segundo a tercero"
]

# === Obtener totales de estudiantes y docentes ===
total_estudiantes = st.session_state.get("total_estudiantes", 0)
total_docentes = st.session_state.get("total_docentes", 0)

viñetas_actualizadas = [
    f"• {total_estudiantes} estudiantes",
    f"• {total_docentes} docentes"
]

# === Reemplazos en los párrafos del documento ===
for p in doc.paragraphs:
    # Reemplazo de nombre de la organización
    if "Fundación Santo Domingo" in p.text:
        p.text = p.text.replace("Fundación Santo Domingo", nombre_organizacion)

    # Reemplazo de municipio
    if "Barú" in p.text:
        p.text = p.text.replace("Barú", municipio)

    # Reemplazo de cantidad de sedes (con gramática)
    for texto_original in reemplazos_cantidad:
        if texto_original in p.text:
            p.text = p.text.replace(texto_original, texto_sedes)

    # Reemplazo de nombres de sedes
    if marcador_sedes_fijas in p.text:
        p.text = p.text.replace(marcador_sedes_fijas, sedes_como_texto)

    # Eliminar viñetas antiguas
    for antigua in viñetas_antiguas:
        if antigua in p.text:
            p.text = ""

# === Agregar nuevas viñetas con datos reales ===
for viñeta in viñetas_actualizadas:
    doc.add_paragraph(viñeta)

# === Guardar documento personalizado ===
doc.save(ruta_salida)

# === Mostrar botón para descarga ===
with open(ruta_salida, "rb") as f:
    st.download_button(
        label="📄 Descargar propuesta Word",
        data=f,
        file_name="Propuesta_FormulIA.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
